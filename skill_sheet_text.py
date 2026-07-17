"""スキルシート添付のテキスト抽出（PDF / Word(.docx) / Excel(.xlsx) 対応・例外を外に出さない）。

添付ファイルの形式を内容（ヘッダー）から判定して抽出する。
解析に失敗しても例外にせず、警告メッセージを返して提案登録を続行できるようにする。
"""

from __future__ import annotations

import logging
import os
import zipfile
from io import BytesIO
from typing import Optional, Tuple

logger = logging.getLogger(__name__)


def _build_warning(filename: str, reason: str) -> str:
    name = os.path.basename(filename or "") or "添付ファイル"
    return f"スキルシート「{name}」を読み込めませんでした（{reason}）。提案はスキルシート本文なしで登録しました。"


def _detect_format(filename: str, data: bytes) -> Optional[str]:
    """先頭バイトと拡張子からファイル形式を判定する。

    Args:
        filename: 添付ファイル名（拡張子判定のフォールバックに使用）。
        data: ファイルの全バイト列。

    Returns:
        Optional[str]: "pdf" / "docx" / "xlsx"。判定不能なら None。
    """
    if data.startswith(b"%PDF"):
        return "pdf"
    if data.startswith(b"PK"):
        try:
            with zipfile.ZipFile(BytesIO(data)) as zf:
                names = set(zf.namelist())
        except zipfile.BadZipFile:
            return None
        if any(n.startswith("word/") for n in names):
            return "docx"
        if any(n.startswith("xl/") for n in names):
            return "xlsx"
        return None
    ext = os.path.splitext(filename or "")[1].lower()
    if ext == ".pdf":
        return "pdf"
    return None


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(BytesIO(data))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _extract_docx(data: bytes) -> str:
    from docx import Document

    document = Document(BytesIO(data))
    parts = [p.text for p in document.paragraphs if p.text]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append("\t".join(cells))
    return "\n".join(parts)


def _extract_xlsx(data: bytes) -> str:
    from openpyxl import load_workbook

    workbook = load_workbook(BytesIO(data), read_only=True, data_only=True)
    parts = []
    try:
        for sheet in workbook.worksheets:
            parts.append(f"[シート: {sheet.title}]")
            for row in sheet.iter_rows(values_only=True):
                cells = [str(v).strip() for v in row if v is not None and str(v).strip()]
                if cells:
                    parts.append("\t".join(cells))
    finally:
        workbook.close()
    return "\n".join(parts)


def extract_skill_sheet_text_safe(skill_sheet) -> Tuple[str, Optional[str]]:
    """スキルシートからテキストを抽出する（失敗しても例外を送出しない）。

    Args:
        skill_sheet: UploadFile 互換オブジェクト（.filename / .file または read 可能）。

    Returns:
        Tuple[str, Optional[str]]: (抽出テキスト, 警告メッセージ)。
            成功時は (テキスト, None)、失敗時は ("", 警告文言)。
    """
    if not skill_sheet:
        return "", None

    filename = getattr(skill_sheet, "filename", None) or getattr(skill_sheet, "name", "") or ""
    file_obj = getattr(skill_sheet, "file", None) or skill_sheet
    if not hasattr(file_obj, "read"):
        return "", _build_warning(filename, "ファイルとして読み取れません")

    current_pos = None
    try:
        if hasattr(file_obj, "tell") and hasattr(file_obj, "seek"):
            current_pos = file_obj.tell()
            file_obj.seek(0)
        data = file_obj.read()
    except Exception as exc:
        logger.warning("スキルシートの読み込みに失敗しました: %s", exc)
        return "", _build_warning(filename, "読み込みに失敗しました")
    finally:
        if current_pos is not None:
            try:
                file_obj.seek(current_pos)
            except Exception:
                pass

    if not data:
        return "", _build_warning(filename, "ファイルが空です")

    kind = _detect_format(filename, data)
    if kind is None:
        logger.warning("スキルシートの形式を判定できません: filename=%s header=%s", filename, data[:8].hex())
        return "", _build_warning(filename, "対応していない形式です。PDF・Word(.docx)・Excel(.xlsx)を添付してください")

    try:
        if kind == "pdf":
            return _extract_pdf(data), None
        if kind == "docx":
            return _extract_docx(data), None
        return _extract_xlsx(data), None
    except Exception as exc:
        logger.warning("スキルシート(%s)の解析に失敗しました: %s", kind, exc)
        return "", _build_warning(filename, "ファイルの解析に失敗しました")
